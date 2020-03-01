import datetime
import uuid

import qrcode
import win32api
import win32print
from docx import Document
from docx.shared import Inches


def createDOCX(filename, rows, cols):
    document = Document()
    document.add_heading("样本盘二维码", level=1)
    document.add_heading("日期：" + str(datetime.date.today()), level=2)
    table = document.add_table(rows=rows, cols=cols)
    imagefile = "d:/_tmp.png"
    for i in range(rows):
        for j in range(cols):
            randomStr = ''.join(str(uuid.uuid4()).split('-'))
            image = qrcode.make(randomStr + '刘鹏').save(imagefile)
            numStr = '第' + str(i) + '行第' + str(j) + '列'
            table.cell(i, j).add_paragraph(randomStr)
            paragraph = table.cell(i, j).add_paragraph(numStr)
            run = paragraph.add_run()
            run.add_picture(imagefile, width=Inches(1))
    document.save(filename)


def printDOCX(filename):
    open(filename, "r")
    win32api.ShellExecute(0, "print", filename, '/d:"%s"' % win32print.GetDefaultPrinter(), ".", 0)


if __name__ == '__main__':
    filename = "d:/test.docx"
    rows = 4
    cols = 5
    createDOCX(filename=filename, rows=rows, cols=cols)
    printDOCX(filename=filename)
